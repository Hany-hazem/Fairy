"""
File Content Analysis Engine

This module provides content analysis capabilities for various file formats including
text extraction, PDF processing, image OCR, and document structure analysis.
"""

import re
import logging
import mimetypes
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple, Set
from dataclasses import dataclass, field
from enum import Enum
import hashlib
import json

# Optional imports for advanced features
try:
    import PyPDF2
    import pdfplumber
    HAS_PDF_SUPPORT = True
except ImportError:
    HAS_PDF_SUPPORT = False

try:
    from PIL import Image
    import pytesseract
    HAS_OCR_SUPPORT = True
except ImportError:
    HAS_OCR_SUPPORT = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX_SUPPORT = True
except ImportError:
    HAS_DOCX_SUPPORT = False

try:
    import openpyxl
    import pandas as pd
    HAS_EXCEL_SUPPORT = True
except ImportError:
    HAS_EXCEL_SUPPORT = False


class ContentType(Enum):
    """Types of content that can be analyzed"""
    TEXT = "text"
    PDF = "pdf"
    IMAGE = "image"
    DOCUMENT = "document"
    SPREADSHEET = "spreadsheet"
    CODE = "code"
    MARKUP = "markup"
    DATA = "data"
    BINARY = "binary"
    UNKNOWN = "unknown"


@dataclass
class ExtractedEntity:
    """Extracted entity from content"""
    text: str
    entity_type: str
    confidence: float
    start_pos: int
    end_pos: int
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentStructure:
    """Document structure analysis"""
    headings: List[Dict[str, Any]] = field(default_factory=list)
    paragraphs: List[str] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    links: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ContentAnalysis:
    """Complete content analysis result"""
    file_path: str
    content_type: ContentType
    extracted_text: str
    original_content: str
    file_size: int
    language: Optional[str] = None
    encoding: str = "utf-8"
    entities: List[ExtractedEntity] = field(default_factory=list)
    structure: Optional[DocumentStructure] = None
    summary: Optional[str] = None
    keywords: List[str] = field(default_factory=list)
    topics: List[str] = field(default_factory=list)
    sentiment: Optional[Dict[str, float]] = None
    readability_score: Optional[float] = None
    word_count: int = 0
    char_count: int = 0
    line_count: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)
    analysis_timestamp: datetime = field(default_factory=datetime.now)
    processing_time: float = 0.0


class FileContentAnalyzer:
    """Analyzes file content and extracts meaningful information"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Content type mappings
        self.content_type_map = {
            '.txt': ContentType.TEXT,
            '.md': ContentType.MARKUP,
            '.html': ContentType.MARKUP,
            '.htm': ContentType.MARKUP,
            '.xml': ContentType.MARKUP,
            '.json': ContentType.DATA,
            '.csv': ContentType.DATA,
            '.yaml': ContentType.DATA,
            '.yml': ContentType.DATA,
            '.pdf': ContentType.PDF,
            '.docx': ContentType.DOCUMENT,
            '.doc': ContentType.DOCUMENT,
            '.rtf': ContentType.DOCUMENT,
            '.odt': ContentType.DOCUMENT,
            '.xlsx': ContentType.SPREADSHEET,
            '.xls': ContentType.SPREADSHEET,
            '.ods': ContentType.SPREADSHEET,
            '.py': ContentType.CODE,
            '.js': ContentType.CODE,
            '.ts': ContentType.CODE,
            '.java': ContentType.CODE,
            '.cpp': ContentType.CODE,
            '.c': ContentType.CODE,
            '.h': ContentType.CODE,
            '.css': ContentType.CODE,
            '.sql': ContentType.CODE,
            '.jpg': ContentType.IMAGE,
            '.jpeg': ContentType.IMAGE,
            '.png': ContentType.IMAGE,
            '.gif': ContentType.IMAGE,
            '.bmp': ContentType.IMAGE,
            '.tiff': ContentType.IMAGE,
            '.webp': ContentType.IMAGE,
        }
        
        # Entity patterns for basic entity extraction
        self.entity_patterns = {
            'email': re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
            'phone': re.compile(r'\b(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b'),
            'url': re.compile(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'),
            'date': re.compile(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b'),
            'time': re.compile(r'\b\d{1,2}:\d{2}(?::\d{2})?\s?(?:AM|PM|am|pm)?\b'),
            'currency': re.compile(r'\$\d+(?:,\d{3})*(?:\.\d{2})?|\b\d+(?:,\d{3})*(?:\.\d{2})?\s?(?:USD|EUR|GBP|dollars?|euros?|pounds?)\b'),
            'ssn': re.compile(r'\b\d{3}-\d{2}-\d{4}\b'),
            'credit_card': re.compile(r'\b\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}\b'),
        }
    
    async def analyze_file_content(self, file_path: str, content: str) -> ContentAnalysis:
        """Analyze file content and extract information"""
        start_time = datetime.now()
        
        try:
            path = Path(file_path)
            content_type = self._determine_content_type(path)
            
            # Initialize analysis result
            analysis = ContentAnalysis(
                file_path=file_path,
                content_type=content_type,
                extracted_text="",
                original_content=content,
                file_size=len(content.encode('utf-8'))
            )
            
            # Extract text based on content type
            if content_type == ContentType.PDF:
                analysis.extracted_text = await self._extract_pdf_text(file_path)
            elif content_type == ContentType.IMAGE:
                analysis.extracted_text = await self._extract_image_text(file_path)
            elif content_type == ContentType.DOCUMENT:
                analysis.extracted_text = await self._extract_document_text(file_path)
            elif content_type == ContentType.SPREADSHEET:
                analysis.extracted_text = await self._extract_spreadsheet_text(file_path)
            else:
                analysis.extracted_text = content
            
            # Perform text analysis
            await self._analyze_text_content(analysis)
            
            # Extract entities
            analysis.entities = await self._extract_entities(analysis.extracted_text)
            
            # Analyze document structure
            if content_type in [ContentType.MARKUP, ContentType.DOCUMENT]:
                analysis.structure = await self._analyze_document_structure(content, content_type)
            
            # Calculate processing time
            end_time = datetime.now()
            analysis.processing_time = (end_time - start_time).total_seconds()
            
            return analysis
            
        except Exception as e:
            self.logger.error(f"Error analyzing file content {file_path}: {e}")
            # Return basic analysis on error
            return ContentAnalysis(
                file_path=file_path,
                content_type=ContentType.UNKNOWN,
                extracted_text=content[:1000] if content else "",
                original_content=content,
                file_size=len(content.encode('utf-8')) if content else 0,
                metadata={'error': str(e)}
            )
    
    def _determine_content_type(self, path: Path) -> ContentType:
        """Determine content type from file extension"""
        extension = path.suffix.lower()
        return self.content_type_map.get(extension, ContentType.UNKNOWN)
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        if not HAS_PDF_SUPPORT:
            self.logger.warning("PDF support not available. Install PyPDF2 and pdfplumber.")
            return ""
        
        try:
            text = ""
            
            # Try with pdfplumber first (better for complex layouts)
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception:
                # Fallback to PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Error extracting PDF text from {file_path}: {e}")
            return ""
    
    async def _extract_image_text(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        if not HAS_OCR_SUPPORT:
            self.logger.warning("OCR support not available. Install PIL and pytesseract.")
            return ""
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Error extracting text from image {file_path}: {e}")
            return ""
    
    async def _extract_document_text(self, file_path: str) -> str:
        """Extract text from document files (DOCX, etc.)"""
        path = Path(file_path)
        
        if path.suffix.lower() == '.docx' and HAS_DOCX_SUPPORT:
            try:
                doc = DocxDocument(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text.strip()
            except Exception as e:
                self.logger.error(f"Error extracting DOCX text from {file_path}: {e}")
        
        return ""
    
    async def _extract_spreadsheet_text(self, file_path: str) -> str:
        """Extract text from spreadsheet files"""
        if not HAS_EXCEL_SUPPORT:
            self.logger.warning("Excel support not available. Install openpyxl and pandas.")
            return ""
        
        try:
            path = Path(file_path)
            
            if path.suffix.lower() in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path, sheet_name=None)
                text = ""
                
                for sheet_name, sheet_df in df.items():
                    text += f"Sheet: {sheet_name}\n"
                    text += sheet_df.to_string() + "\n\n"
                
                return text.strip()
            
        except Exception as e:
            self.logger.error(f"Error extracting spreadsheet text from {file_path}: {e}")
        
        return ""
    
    async def _analyze_text_content(self, analysis: ContentAnalysis):
        """Analyze text content for basic metrics"""
        text = analysis.extracted_text
        
        if not text:
            return
        
        # Basic text metrics
        analysis.word_count = len(text.split())
        analysis.char_count = len(text)
        analysis.line_count = len(text.splitlines())
        
        # Extract keywords (simple approach - most frequent words)
        words = re.findall(r'\b\w+\b', text.lower())
        word_freq = {}
        for word in words:
            if len(word) > 3:  # Skip short words
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Get top keywords
        analysis.keywords = sorted(word_freq.keys(), key=word_freq.get, reverse=True)[:20]
        
        # Simple topic extraction (based on keywords)
        analysis.topics = self._extract_topics(analysis.keywords)
        
        # Language detection (simple heuristic)
        analysis.language = self._detect_language(text)
        
        # Generate simple summary
        analysis.summary = self._generate_summary(text)
    
    async def _extract_entities(self, text: str) -> List[ExtractedEntity]:
        """Extract entities from text using regex patterns"""
        entities = []
        
        for entity_type, pattern in self.entity_patterns.items():
            for match in pattern.finditer(text):
                entity = ExtractedEntity(
                    text=match.group(),
                    entity_type=entity_type,
                    confidence=0.8,  # Simple confidence score
                    start_pos=match.start(),
                    end_pos=match.end()
                )
                entities.append(entity)
        
        return entities
    
    async def _analyze_document_structure(self, content: str, content_type: ContentType) -> DocumentStructure:
        """Analyze document structure"""
        structure = DocumentStructure()
        
        if content_type == ContentType.MARKUP:
            # HTML/Markdown structure analysis
            structure.headings = self._extract_headings(content)
            structure.links = self._extract_links(content)
            structure.paragraphs = self._extract_paragraphs(content)
        
        return structure
    
    def _extract_headings(self, content: str) -> List[Dict[str, Any]]:
        """Extract headings from markup content"""
        headings = []
        
        # HTML headings
        html_heading_pattern = re.compile(r'<h([1-6])[^>]*>(.*?)</h\1>', re.IGNORECASE | re.DOTALL)
        for match in html_heading_pattern.finditer(content):
            headings.append({
                'level': int(match.group(1)),
                'text': re.sub(r'<[^>]+>', '', match.group(2)).strip(),
                'start_pos': match.start(),
                'end_pos': match.end()
            })
        
        # Markdown headings
        md_heading_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
        for match in md_heading_pattern.finditer(content):
            headings.append({
                'level': len(match.group(1)),
                'text': match.group(2).strip(),
                'start_pos': match.start(),
                'end_pos': match.end()
            })
        
        return headings
    
    def _extract_links(self, content: str) -> List[Dict[str, Any]]:
        """Extract links from markup content"""
        links = []
        
        # HTML links
        html_link_pattern = re.compile(r'<a[^>]+href=["\']([^"\']+)["\'][^>]*>(.*?)</a>', re.IGNORECASE | re.DOTALL)
        for match in html_link_pattern.finditer(content):
            links.append({
                'url': match.group(1),
                'text': re.sub(r'<[^>]+>', '', match.group(2)).strip(),
                'start_pos': match.start(),
                'end_pos': match.end()
            })
        
        # Markdown links
        md_link_pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
        for match in md_link_pattern.finditer(content):
            links.append({
                'text': match.group(1),
                'url': match.group(2),
                'start_pos': match.start(),
                'end_pos': match.end()
            })
        
        return links
    
    def _extract_paragraphs(self, content: str) -> List[str]:
        """Extract paragraphs from content"""
        # Simple paragraph extraction
        paragraphs = []
        
        # Split by double newlines
        raw_paragraphs = re.split(r'\n\s*\n', content)
        
        for para in raw_paragraphs:
            # Clean up paragraph
            clean_para = re.sub(r'<[^>]+>', '', para)  # Remove HTML tags
            clean_para = re.sub(r'\s+', ' ', clean_para)  # Normalize whitespace
            clean_para = clean_para.strip()
            
            if len(clean_para) > 20:  # Only include substantial paragraphs
                paragraphs.append(clean_para)
        
        return paragraphs
    
    def _extract_topics(self, keywords: List[str]) -> List[str]:
        """Extract topics from keywords (simple approach)"""
        # This is a simplified topic extraction
        # In a real implementation, you might use more sophisticated NLP techniques
        
        topic_keywords = {
            'technology': ['software', 'computer', 'system', 'data', 'code', 'programming', 'development'],
            'business': ['company', 'market', 'sales', 'revenue', 'customer', 'business', 'strategy'],
            'science': ['research', 'study', 'analysis', 'experiment', 'theory', 'method', 'results'],
            'education': ['learning', 'student', 'course', 'education', 'teaching', 'knowledge', 'skill'],
            'health': ['health', 'medical', 'patient', 'treatment', 'disease', 'medicine', 'care'],
        }
        
        topics = []
        for topic, topic_words in topic_keywords.items():
            if any(keyword in topic_words for keyword in keywords[:10]):  # Check top keywords
                topics.append(topic)
        
        return topics
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection (placeholder)"""
        # This is a very basic implementation
        # In practice, you'd use a proper language detection library
        
        if not text:
            return "unknown"
        
        # Simple heuristic based on common words
        english_words = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
        words = set(text.lower().split())
        
        english_score = len(words.intersection(english_words))
        
        if english_score > 0:
            return "en"
        
        return "unknown"
    
    def _generate_summary(self, text: str, max_sentences: int = 3) -> str:
        """Generate a simple summary of the text"""
        if not text or len(text) < 100:
            return text[:200] + "..." if len(text) > 200 else text
        
        # Simple extractive summarization
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
        
        if len(sentences) <= max_sentences:
            return '. '.join(sentences) + '.'
        
        # Take first, middle, and last sentences as a simple summary
        summary_sentences = [
            sentences[0],
            sentences[len(sentences) // 2],
            sentences[-1]
        ]
        
        return '. '.join(summary_sentences) + '.'